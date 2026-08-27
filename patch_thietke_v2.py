"""
Script auto-patch ThietKeHeThong_AILIP_PERFECT.docx -> ThietKeHeThong_AILIP_FINAL.docx
"""
import copy, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_PATH  = r'd:\NCKH\ThietKeHeThong_AILIP_PERFECT.docx'
OUTPUT_PATH = r'd:\NCKH\ThietKeHeThong_AILIP_FINAL.docx'

doc = Document(INPUT_PATH)

def find_para(doc, keyword):
    for p in doc.paragraphs:
        if keyword in p.text:
            return p
    return None

def insert_para_after(ref, text, style='List Paragraph'):
    new_elem = OxmlElement('w:p')
    ref._element.addnext(new_elem)
    from docx.text.paragraph import Paragraph as P
    new_p = None
    for p in doc.paragraphs:
        if p._element is new_elem:
            new_p = p
            break
    if new_p is None:
        new_p = P(new_elem, ref._element.getparent())
    try:
        new_p.style = doc.styles[style]
    except:
        pass
    new_p.add_run(text)
    return new_p

# ── 1. Vênh phạm vi (Scope Notes) ─────────────────────────────
# Thêm ghi chú cho organizations
p_org = find_para(doc, 'organizations')
if p_org:
    note = "Ghi ch\u00fa: B\u1ea3ng organizations \u0111\u01b0\u1ee3c thi\u1ebft k\u1ebf s\u1eb5n \u0111\u1ec3 \u0111\u1ea3m b\u1ea3o kh\u1ea3 n\u0103ng m\u1edf r\u1ed9ng (forward-compatible), nh\u01b0ng ch\u1ee9c n\u0103ng doanh nghi\u1ec7p n\u00e2ng cao t\u1ea1m ho\u00e3n trong ph\u1ea1m vi 10 tu\u1ea7n c\u1ee7a \u0111\u1ed3 \u00e1n."
    insert_para_after(p_org, note, style='Normal').runs[0].italic = True

# Thêm ghi chú cho collections
p_col = find_para(doc, 'collections')
if p_col:
    note = "Ghi ch\u00fa: Tr\u01b0\u1eddng is_shared (n\u1ebfu c\u00f3) \u0111\u01b0\u1ee3c thi\u1ebft k\u1ebf d\u1ef1 ph\u00f2ng cho t\u00ednh n\u0103ng chia s\u1ebb nh\u00f3m nh\u01b0ng ch\u01b0a tri\u1ec3n khai \u1edf phi\u00ean b\u1ea3n hi\u1ec7n t\u1ea1i."
    insert_para_after(p_col, note, style='Normal').runs[0].italic = True

# Thêm ghi chú cho RBAC
p_rbac = find_para(doc, 'Ph\u00e2n quy\u1ec1n (RBAC')
if p_rbac:
    note = "Ghi ch\u00fa: Vai tr\u00f2 Enterprise User \u0111\u01b0\u1ee3c \u0111\u1ecbnh ngh\u0129a \u0111\u1ec3 s\u1eb5n s\u00e0ng \u0111\u1ed3ng b\u1ed9 v\u1edbi b\u1ea3ng organizations, tuy nhi\u00ean s\u1ebd kh\u00f4ng ho\u1ea1t \u0111\u1ed9ng trong ph\u1ea1m vi 10 tu\u1ea7n (out-of-scope)."
    insert_para_after(p_rbac, note, style='Normal').runs[0].italic = True

print("[OK] 1. Fixed scope notes (v\u00eanh ph\u1ea1m vi)")


# ── 2. Ma trận truy vết (Traceability Matrix) ─────────────────
# Tìm 1.3 Component Diagram để chèn vào cuối chương 1
def find_end_of_chapter(doc, start_heading):
    found = False
    last_p = None
    for p in doc.paragraphs:
        if start_heading in p.text and p.style and p.style.name.startswith('Heading'):
            found = True
        if found:
            if p.style and p.style.name.startswith('Heading 1'):
                if start_heading not in p.text:
                    return last_p # End of current chapter
            last_p = p
    return last_p

ch1_end = find_end_of_chapter(doc, 'CH\u01af\u01a0NG 1')
if ch1_end:
    h = insert_para_after(ch1_end, "1.4. Ma tr\u1eadn Truy v\u1ebft (Traceability Matrix)", style='Heading 2')
    h.runs[0].bold = True
    desc = insert_para_after(h, "B\u1ea3ng d\u01b0\u1edbi \u0111\u00e2y \u0111\u1ed1i chi\u1ebfu c\u00e1c Use Case t\u1eeb Ph\u00e2n t\u00edch H\u1ec7 th\u1ed1ng \u0111\u1ec3 ch\u1ee9ng minh Thi\u1ebft k\u1ebf \u0111\u00e3 ph\u1ee7 \u0111\u1ea7y \u0111\u1ee7 v\u00e0 kh\u00f4ng \u0111i l\u1ec7ch kh\u1ecfi ph\u1ea1m vi:", style='Normal')
    
    matrix_data = [
        ("Nhu c\u1ea7u / Use Case", "B\u1ea3ng D\u1eef li\u1ec7u", "API Endpoint", "M\u00e0n h\u00ecnh (UI)"),
        ("Tra c\u1ee9u v\u0103n b\u1ea3n l\u1edbn", "documents, document_chunks", "GET /api/v1/search", "SearchPage"),
        ("H\u1ecfi \u0111\u00e1p AI ph\u00e1p l\u00fd (RAG)", "documents, document_chunks", "POST /api/v1/ai/chat", "Chat AI"),
        ("Xem Knowledge Graph", "document_relations", "GET /api/v1/graph/{id}", "Knowledge Graph"),
        ("T\u00f3m t\u1eaft v\u0103n b\u1ea3n m\u1edbi", "documents", "POST /api/v1/ai/summarize", "DocumentPage"),
        ("Qu\u1ea3n l\u00fd t\u00e0i li\u1ec7u c\u00e1 nh\u00e2n", "collections, notes", "GET /api/v1/workspace", "WorkspacePage"),
        ("Ph\u00e2n t\u00edch xu h\u01b0\u1edbng", "documents", "GET /api/v1/analytics", "AnalyticsPage"),
    ]
    
    table = doc.add_table(rows=len(matrix_data), cols=4)
    # table.style = 'Table Grid'
    for row_idx, row_data in enumerate(matrix_data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            run = cell.paragraphs[0].add_run(cell_text)
            if row_idx == 0:
                run.bold = True
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D9D9D9')
                tcPr.append(shd)
    
    tbl_el = table._tbl
    doc.element.body.remove(tbl_el)
    desc._element.addnext(tbl_el)
    print("[OK] 2. Added Traceability Matrix")


# ── 3. Thiết kế API (/api/v1 và format lỗi) ───────────────────
p_api = find_para(doc, '4.1. Ti\u00eau chu\u1ea9n')
if p_api:
    insert_para_after(p_api, "C\u1ea5u tr\u00fac L\u1ed7i chu\u1ea9n (Standard Error Response): H\u1ec7 th\u1ed1ng th\u1ed1ng nh\u1ea5t tr\u1ea3 v\u1ec1 JSON c\u00f3 d\u1ea1ng {\"error\": {\"code\": \"ERROR_CODE\", \"message\": \"M\u00f4 t\u1ea3\"}} thay v\u00ec text thu\u1ea7n, \u0111\u1ec3 frontend d\u1ec5 x\u1eed l\u00fd.")
    insert_para_after(p_api, "Ti\u1ec1n t\u1ed1 API: To\u00e0n b\u1ed9 endpoint \u0111\u01b0\u1ee3c \u0111\u1eb7t d\u01b0\u1edbi ti\u1ec1n t\u1ed1 /api/v1/ \u0111\u1ec3 \u0111\u1ea3m b\u1ea3o an to\u00e0n khi n\u00e2ng c\u1ea5p (API Versioning).")

# Update existing table cells
for table in doc.tables:
    if len(table.rows) > 0 and len(table.columns) >= 2:
        if 'Endpoint' in table.rows[0].cells[1].text:
            # Change "Auth" to "Role"
            if len(table.columns) >= 3 and table.rows[0].cells[2].text.strip() == "Auth":
                table.rows[0].cells[2].paragraphs[0].runs[0].text = "Vai tr\u00f2 (Role)"
            
            for row in table.rows[1:]:
                # Thêm /api/v1
                endpoint_text = row.cells[1].text.strip()
                if endpoint_text.startswith('/'):
                    for p in row.cells[1].paragraphs:
                        p.text = "/api/v1" + p.text
                
                # Role
                if len(row.cells) >= 3:
                    auth_text = row.cells[2].text.strip()
                    for p in row.cells[2].paragraphs:
                        if "Khong" in auth_text or "Không" in auth_text:
                            p.text = "Public"
                        elif "Bearer" in auth_text:
                            p.text = "User/Admin"
                            
            print("[OK] 3. Updated API versions and error format")
            break

# ── 4. Yêu cầu phi chức năng (NFR) ───────────────────────────
ch7_end = find_end_of_chapter(doc, 'CH\u01af\u01a0NG 7')
if ch7_end:
    h = insert_para_after(ch7_end, "CH\u01af\u01a0NG 8: Y\u00caU C\u1ea6U PHI CH\u1ee8C N\u0102NG (NON-FUNCTIONAL REQUIREMENTS)", style='Heading 1')
    
    nfrs = [
        "Th\u1eddi gian ph\u1ea3n h\u1ed3i (Latency): Truy v\u1ea5n t\u00ecm ki\u1ebfm th\u00f4ng th\u01b0\u1eddng (BM25/Semantic) ph\u1ea3i ho\u00e0n th\u00e0nh v\u00e0 tr\u1ea3 v\u1ec1 d\u01b0\u1edbi 2 gi\u00e2y.",
        "AI First-Token Latency: T\u1eeb l\u00fac ng\u01b0\u1eddi d\u00f9ng g\u1eedi c\u00e2u h\u1ecfi ph\u00e1p l\u00fd, AI ph\u1ea3i b\u1eaft \u0111\u1ea7u stream k\u00fd t\u1ef1 \u0111\u1ea7u ti\u00ean l\u00ean m\u00e0n h\u00ecnh d\u01b0\u1edbi 3 gi\u00e2y \u0111\u1ec3 \u0111\u1ea3m b\u1ea3o UX (L\u00fd do thi\u1ebft l\u1eadp timeout 10s c\u1ee7a Circuit Breaker).",
        "S\u1ee9c ch\u1ecbu t\u1ea3i (Concurrency): H\u1ed7 tr\u1ee3 t\u1ed1i thi\u1ec3u 50 users thao t\u00e1c \u0111\u1ed3ng th\u1eddi tr\u00ean ki\u1ebfn tr\u00fac Docker m\u1ed9t m\u00e1y ch\u1ee7 (8GB RAM), s\u1eed d\u1ee5ng FastAPI async workers.",
    ]
    ref = h
    for nfr in nfrs:
        ref = insert_para_after(ref, nfr)
    print("[OK] 4. Added NFR chapter")


# ── 5. Rate Limiting ──────────────────────────────────────────
for para in doc.paragraphs:
    if 'N request' in para.text:
        for r in para.runs:
            if 'N request/ph\u00fat' in r.text or 'N request' in r.text:
                r.text = r.text.replace('N request/ph\u00fat', '15 request/ph\u00fat/user').replace('N request', '15 request/ph\u00fat/user')
        print("[OK] 5. Fixed Rate limiting numbers")
        break


# ── 6. Tổng kết thiết kế ──────────────────────────────────────
last_p = doc.paragraphs[-1]
h = insert_para_after(last_p, "K\u1ebeT LU\u1eacN THI\u1ebeT K\u1ebe", style='Heading 1')
content = "T\u00e0i li\u1ec7u thi\u1ebft k\u1ebf n\u00e0y \u0111\u00e3 v\u1eadt l\u00fd h\u00f3a ho\u00e0n to\u00e0n 6 nhu c\u1ea7u c\u1ed1t l\u00f5i \u0111\u01b0\u1ee3c v\u1ea1ch ra \u1edf giai \u0111o\u1ea1n ph\u00e2n t\u00edch (Tra c\u1ee9u v\u0103n b\u1ea3n, H\u1ecfi \u0111\u00e1p AI, T\u00f3m t\u1eaft, Dashboard, Workspace, Knowledge Graph). Vi\u1ec7c \u0111\u1eb7t r\u00f5 r\u00e0ng ma tr\u1eadn truy v\u1ebft c\u0169ng nh\u01b0 gi\u1edbi h\u1ea1n \u0111\u1ec3 ng\u1ecf c\u00e1c th\u1ef1c th\u1ec3 doanh nghi\u1ec7p cho t\u01b0\u01a1ng lai \u0111\u1ea3m b\u1ea3o h\u1ec7 th\u1ed1ng th\u1ecfa m\u00e3n \u0111\u00fang ph\u1ea1m vi 10 tu\u1ea7n c\u1ee7a \u0111\u1ed3 \u00e1n, tr\u00e1nh \u0111\u01b0\u1ee3c r\u1ee7i ro 'ph\u00ecnh ph\u1ea1m vi' (scope creep). Giai \u0111o\u1ea1n ti\u1ebfp theo s\u1ebd ti\u1ebfn h\u00e0nh l\u1eadp tr\u00ecnh d\u1ef1a tr\u00ean class/service architecture n\u00e0y."
insert_para_after(h, content, style='Normal')
print("[OK] 6. Added Design Conclusion")

doc.save(OUTPUT_PATH)
print(f"\n✅ All 6 advanced academic fixes applied. Saved to: {OUTPUT_PATH}")
