import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc_path = 'd:/NCKH/PhanTichHeThong_v2_Fixed_Updated.docx'
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # Medium Grid 1 Accent 1 is a clean blue-ish grid in modern MS Word
    target_style = 'Medium Grid 1 Accent 1'
    
    for i, table in enumerate(doc.tables):
        try:
            table.style = target_style
        except Exception:
            try:
                table.style = 'Table Grid'
            except:
                pass
        
        # Style Header
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)

        # Style content
        for row in table.rows[1:]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    
    # Color Headings
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 82, 155)  # Professional Dark Blue
    
    try:
        doc.save(doc_path)
        print("Success")
    except Exception as e:
        print(f"Save error: {e}")

if __name__ == '__main__':
    main()
