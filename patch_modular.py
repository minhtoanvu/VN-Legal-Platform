import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement

DOC_PATH = r'd:\NCKH\ThietKeHeThong_AILIP_FINAL.docx'
doc = Document(DOC_PATH)

def insert_para_after(ref, text, style='Normal', bold_text=None):
    new_elem = OxmlElement('w:p')
    ref._element.addnext(new_elem)
    from docx.text.paragraph import Paragraph as P
    new_p = None
    for p in doc.paragraphs:
        if p._element is new_elem:
            new_p = p
            break
    if new_p is None:
        new_p = P(new_elem, ref._element.getparent())
    try:
        new_p.style = doc.styles[style]
    except:
        pass
    
    if bold_text and bold_text in text:
        parts = text.split(bold_text)
        new_p.add_run(parts[0])
        r = new_p.add_run(bold_text)
        r.bold = True
        if len(parts) > 1:
            new_p.add_run(parts[1])
    else:
        new_p.add_run(text)
    return new_p

found = False
for p in doc.paragraphs:
    if '1.2' in p.text and 'Ki\u1ebfn tr\u00fac' in p.text:
        note = "H\u1ec7 th\u1ed1ng API Server \u0111\u01b0\u1ee3c thi\u1ebft k\u1ebf theo ki\u1ebfn tr\u00fac Modular Monolith (Nguy\u00ean kh\u1ed1i theo M\u00f4-\u0111un). M\u1eb7c d\u00f9 to\u00e0n b\u1ed9 Backend ch\u1ea1y tr\u00ean m\u1ed9t service duy nh\u1ea5t \u0111\u1ec3 t\u1ed1i \u01b0u chi ph\u00ed h\u1ea1 t\u1ea7ng (ph\u00f9 h\u1ee3p v\u1edbi \u0111\u1ed3 \u00e1n 10 tu\u1ea7n), nh\u01b0ng \u1edf t\u1ea7ng code, c\u00e1c router v\u00e0 service (Auth, RAG, Search) \u0111\u01b0\u1ee3c c\u00f4 l\u1eadp ranh gi\u1edbi r\u00f5 r\u00e0ng. \u0110i\u1ec1u n\u00e0y \u0111\u1ea3m b\u1ea3o t\u00ednh forward-compatible, s\u1eb5n s\u00e0ng t\u00e1ch th\u00e0nh Microservices khi h\u1ec7 th\u1ed1ng scale l\u00ean quy m\u00f4 doanh nghi\u1ec7p m\u00e0 kh\u00f4ng c\u1ea7n \u0111\u1eadp \u0111i x\u00e2y l\u1ea1i."
        insert_para_after(p, note, style='Normal', bold_text='Modular Monolith')
        found = True
        break

if found:
    doc.save(DOC_PATH)
    print("✅ Đã cập nhật kiến trúc Modular Monolith vào file FINAL.")
else:
    print("❌ Không tìm thấy mục 1.2 Kiến trúc.")
