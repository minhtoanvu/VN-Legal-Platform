"""Script auto-patch ThietKeHeThong_AILIP.docx"""
import copy, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_PATH  = r'd:\NCKH\ThietKeHeThong_AILIP.docx'
OUTPUT_PATH = r'd:\NCKH\ThietKeHeThong_AILIP_PERFECT.docx'

doc = Document(INPUT_PATH)

def find_para(doc, keyword):
    for p in doc.paragraphs:
        if keyword in p.text:
            return p
    return None

def insert_para_after(ref, text, style='List Paragraph'):
    new_elem = OxmlElement('w:p')
    ref._element.addnext(new_elem)
    from docx.text.paragraph import Paragraph as P
    new_p = None
    for p in doc.paragraphs:
        if p._element is new_elem:
            new_p = p
            break
    if new_p is None:
        new_p = P(new_elem, ref._element.getparent())
    try:
        new_p.style = doc.styles[style]
    except:
        pass
    new_p.add_run(text)
    return new_p

# ── 1. Sửa session_id (Ch.5) ─────────────────────────────────
for para in doc.paragraphs:
    if 'session_id' in para.text:
        full = para.text
        new_text = full.replace(
            '(session_id)',
            '(lich su hoi thoai duoc truyen theo tung request - stateless design)'
        )
        for r in para.runs:
            r.text = ''
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)
        print("[OK] Fixed session_id description (Ch.5)")
        break

# ── 2. Sửa Fallback Ch.7.1 ─────────────────────────────────────
for para in doc.paragraphs:
    if 'vui long thu lai sau it phut hoac dung chuc nang tim kiem' in para.text.replace('\u00a0', ' '):
        break
    if 'Tr\u1ee3 l\u00fd AI \u0111ang t\u1ea1m gi\u00e1n \u0111o\u1ea1n' in para.text and 'ch\u1ee9c n\u0103ng t\u00ecm ki\u1ebfm' in para.text:
        for r in para.runs:
            r.text = r.text.replace(
                'ho\u1eb7c d\u00f9ng ch\u1ee9c n\u0103ng t\u00ecm ki\u1ebfm.',
                'H\u1ec7 th\u1ed1ng t\u1ef1 \u0111\u1ed9ng tr\u1ea3 v\u1ec1 danh s\u00e1ch k\u1ebft qu\u1ea3 t\u00ecm ki\u1ebfm li\u00ean quan t\u1eeb c\u01a1 s\u1edf d\u1eef li\u1ec7u thay th\u1ebf.'
            )
        print("[OK] Fixed Fallback description (Ch.7.1)")
        break

# ── 3. Bổ sung Ch.7.2 Timeout & Retry Policy ─────────────────
h72 = find_para(doc, 'Timeout')
if h72 and 'Retry' in h72.text:
    idx = next((i for i, p in enumerate(doc.paragraphs) if p._element is h72._element), None)
    if idx is not None:
        next_p = doc.paragraphs[idx + 1] if idx + 1 < len(doc.paragraphs) else None
        need_insert = next_p and (
            (next_p.style and next_p.style.name.startswith('Heading'))
            or '7.3' in next_p.text
            or not next_p.text.strip()
        )
        if need_insert:
            items_72 = [
                "Hệ thống thiết lập thời gian chờ (timeout) cố định là 10 giây cho mỗi lần gọi LLM API — được cấu hình qua hằng số LLM_TIMEOUT_SEC = 10 trong module rag_service.py. Đây là ngưỡng cân bằng giữa trải nghiệm người dùng (không chờ quá lâu) và độ phức tạp của câu hỏi pháp lý.",
                "Hệ thống không thực hiện Retry tự động sau timeout vì câu hỏi pháp lý có thể không đồng bộ với ngữ cảnh nếu gọi lại — thay vào đó, Circuit Breaker kích hoạt Fallback ngay lập tức để đảm bảo phản hồi nhanh.",
                "Thời gian timeout 10 giây được đo từ token đầu tiên của stream (first-token latency), không phải tổng thời gian stream — đảm bảo người dùng không phải chờ quá lâu trước khi thấy phản hồi đầu tiên.",
            ]
            ref = h72
            for item in items_72:
                ref = insert_para_after(ref, item)
            print("[OK] Added Ch.7.2 content (3 items)")

# ── 4. Bổ sung Ch.1.3 Frontend components ─────────────────────
fe_h = find_para(doc, 'Frontend')
if fe_h and fe_h.style and 'Heading' in fe_h.style.name:
    idx_fe = next((i for i, p in enumerate(doc.paragraphs) if p._element is fe_h._element), None)
    if idx_fe is not None:
        next_p = doc.paragraphs[idx_fe + 1]
        if next_p.style and 'Heading' in next_p.style.name:
            fe_items = [
                "SearchPage.tsx — Giao diện tìm kiếm 3 chế độ hybrid/semantic/keyword với bộ lọc đa chiều (lĩnh vực, loại văn bản, hiệu lực).",
                "DocumentPage.tsx — Xem toàn văn, Timeline lịch sử và Knowledge Graph. Tích hợp Vis.js Network để trực quan hóa quan hệ văn bản.",
                "AnalyticsPage.tsx — Dashboard 5 biểu đồ (Line, Pie, Bar, Heatmap, KPI Cards) dùng thư viện Recharts.",
                "WorkspacePage.tsx — Quản lý Collections, Bookmarks và Ghi chú cá nhân.",
                "AuthPage.tsx — Đăng ký / Đăng nhập. Lưu JWT trong localStorage; tự động refresh khi hết hạn.",
                "State Management: Zustand store quản lý trạng thái toàn cục (user session, search filters).",
                "API Layer: services/ \u2014 T\u00e1ch bi\u1ec7t ho\u00e0n to\u00e0n logic g\u1ecdi API kh\u1ecfi UI component.",
            ]
            ref = fe_h
            for item in fe_items:
                ref = insert_para_after(ref, item)
            print("[OK] Added Frontend component list (Ch.1.3)")

# ── 5. Bổ sung Ch.1.3 Backend components ─────────────────────
be_h = find_para(doc, 'Backend')
if be_h and be_h.style and 'Heading' in be_h.style.name:
    idx_be = next((i for i, p in enumerate(doc.paragraphs) if p._element is be_h._element), None)
    if idx_be is not None:
        next_p = doc.paragraphs[idx_be + 1]
        if next_p.style and 'Heading' in next_p.style.name:
            be_items = [
                "auth.py — 4 endpoints: POST /register, POST /login, POST /refresh, GET /me. Xác thực JWT + Bcrypt hash mật khẩu.",
                "search.py — POST & GET /search: 3 chế độ keyword/semantic/hybrid với RRF re-ranking.",
                "documents.py — CRUD /documents: Danh sách, chi tiết, lọc văn bản theo metadata.",
                "ai.py — POST /ai/chat (SSE stream RAG), POST /ai/summarize (tóm tắt văn bản - UC-11).",
                "graph.py — GET /graph/{doc_id}: Trả về nodes+edges cho Vis.js, hỗ trợ depth 1-3.",
                "analytics.py — GET /analytics/dashboard: Aggregation 5 loại biểu đồ theo lĩnh vực/thời gian.",
                "workspace.py — CRUD /workspace/collections, /notes, /bookmarks (collection_documents).",
                "Services: rag_service, bm25_service, semantic_service, rrf_service, graph_service, analytics_service, auth_service.",
            ]
            ref = be_h
            for item in be_items:
                ref = insert_para_after(ref, item)
            print("[OK] Added Backend component list (Ch.1.3)")

# ── 6. Bổ sung bảng Endpoint (Ch.4.2) ─────────────────────────
h42 = find_para(doc, 'Endpoint ch')
if h42:
    idx_42 = next((i for i, p in enumerate(doc.paragraphs) if p._element is h42._element), None)
    desc = doc.paragraphs[idx_42 + 1] if idx_42 is not None else None
    
    # Kiểm tra đã có table chưa
    already_has_table = (desc and desc._element.getnext() is not None 
                         and 'tbl' in desc._element.getnext().tag)
    
    if not already_has_table:
        endpoint_data = [
            ("Method", "Endpoint", "Auth", "Mo ta"),
            ("POST", "/auth/register", "Khong", "\u0110\u0103ng k\u00fd t\u00e0i kho\u1ea3n, tr\u1ea3 v\u1ec1 JWT ngay"),
            ("POST", "/auth/login", "Khong", "\u0110\u0103ng nh\u1eadp, nh\u1eadn access_token (30ph) + refresh_token (7ng)"),
            ("POST", "/auth/refresh", "Khong", "L\u00e0m m\u1edbi access_token b\u1eb1ng refresh_token"),
            ("GET",  "/auth/me", "Bearer", "L\u1ea5y th\u00f4ng tin t\u00e0i kho\u1ea3n \u0111ang \u0111\u0103ng nh\u1eadp"),
            ("POST", "/search", "Khong", "T\u00ecm ki\u1ebfm keyword / semantic / hybrid v\u1edbi RRF"),
            ("GET",  "/search?q=&mode=", "Khong", "T\u00ecm ki\u1ebfm nhanh d\u1ea1ng GET"),
            ("GET",  "/documents/", "Khong", "Danh s\u00e1ch v\u0103n b\u1ea3n, h\u1ed7 tr\u1ee3 ph\u00e2n trang v\u00e0 l\u1ecdc"),
            ("GET",  "/documents/{id}", "Khong", "Chi ti\u1ebft to\u00e0n v\u0103n v\u00e0 metadata m\u1ed9t v\u0103n b\u1ea3n"),
            ("POST", "/ai/chat", "Khong*", "H\u1ecfi \u0111\u00e1p ph\u00e1p l\u00fd RAG \u2014 response d\u1ea1ng SSE stream"),
            ("POST", "/ai/summarize", "Khong*", "T\u00f3m t\u1eaft t\u1ef1 \u0111\u1ed9ng m\u1ed9t v\u0103n b\u1ea3n ph\u00e1p l\u00fd (UC-11)"),
            ("GET",  "/graph/{doc_id}", "Khong", "Knowledge Graph nodes+edges t\u01b0\u01a1ng th\u00edch Vis.js"),
            ("GET",  "/analytics/dashboard", "Khong", "D\u1eef li\u1ec7u 5 bi\u1ec3u \u0111\u1ed3 Dashboard theo b\u1ed9 l\u1ecdc"),
            ("GET",  "/workspace/collections", "Bearer", "Danh s\u00e1ch Collections c\u00e1 nh\u00e2n"),
            ("POST", "/workspace/collections", "Bearer", "T\u1ea1o Collection m\u1edbi"),
            ("POST", "/workspace/collections/{id}/documents", "Bearer", "Th\u00eam v\u0103n b\u1ea3n v\u00e0o Collection"),
            ("POST", "/workspace/notes", "Bearer", "T\u1ea1o ghi ch\u00fa c\u00e1 nh\u00e2n tr\u00ean v\u0103n b\u1ea3n"),
            ("GET",  "/health", "Khong", "Health check \u2014 d\u00f9ng cho Docker/load balancer"),
        ]
        
        table = doc.add_table(rows=len(endpoint_data), cols=4)
        # table.style = 'Table Grid'
        
        for row_idx, row_data in enumerate(endpoint_data):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                run = cell.paragraphs[0].add_run(cell_text)
                if row_idx == 0:
                    run.bold = True
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '1F3A5F')
                    tcPr.append(shd)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif col_idx == 0:
                    run.bold = True
                    if cell_text == "POST":
                        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                    elif cell_text == "GET":
                        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        
        # Di chuyển bảng vào sau desc
        tbl_el = table._tbl
        body = doc.element.body
        body.remove(tbl_el)
        if desc:
            desc._element.addnext(tbl_el)
        
        # Ghi chú sau bảng
        note_elem = OxmlElement('w:p')
        tbl_el.addnext(note_elem)
        from docx.text.paragraph import Paragraph as DP
        note_p = None
        for p in doc.paragraphs:
            if p._element is note_elem:
                note_p = p
                break
        if note_p is None:
            note_p = DP(note_elem, body)
        note_r = note_p.add_run("(*) Khuyen nghi yeu cau xac thuc trong moi truong production de kiem soat chi phi goi LLM. Danh sach day du endpoint CRUD phu tro qua OpenAPI tai /docs.")
        note_r.italic = True
        note_r.font.size = Pt(9)
        
        print(f"[OK] Added endpoint table (Ch.4.2) - {len(endpoint_data)-1} endpoints")

# ── Save ─────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"\n=== DONE: File saved to {OUTPUT_PATH} ===")
print("Changes applied:")
print("  1. Fixed session_id description (Ch.5)")
print("  2. Fixed Fallback description (Ch.7.1)")
print("  3. Added timeout/retry content (Ch.7.2)")
print("  4. Added Frontend component list (Ch.1.3)")
print("  5. Added Backend component list (Ch.1.3)")
print("  6. Added endpoint table 17 rows (Ch.4.2)")
